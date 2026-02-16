import streamlit as st
import json
import fitz  # PyMuPDF
from pathlib import Path
import tempfile
import os
import re
from groq import Groq
from docx import Document  # pip install python-docx
import pandas as pd
import easyocr
import numpy as np
from PIL import Image

# Page configuration
st.set_page_config(
    page_title="Requirement Analysis Tool",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- HELPER FUNCTIONS ---

# def load_config(config_path: str = "config.json") -> dict:
    # try:
    #     with open(config_path, "r", encoding="utf-8") as f:
    #         return json.load(f)
    # except Exception:
    #     return {}
api_key = st.secrets.get("groq_api_key", None)
model_name = st.secrets.get("groq_default_model", "llama-3.3-70b-versatile")

def extract_text_from_file(uploaded_file) -> str:
    """Extracts text from PDF, TXT, DOCX, Excel, Markdown, or Image files."""
    file_extension = Path(uploaded_file.name).suffix.lower()

    try:
        # ---------- IMAGE FILES ----------
        if file_extension in [".png", ".jpg", ".jpeg"]:
            reader = easyocr.Reader(['en'], gpu=False)

            image = Image.open(uploaded_file)
            image_np = np.array(image)

            results = reader.readtext(image_np, detail=0)
            return "\n".join(results)

        # ---------- PDF ----------
        elif file_extension == ".pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name

            doc = fitz.open(tmp_path)
            text = "".join([page.get_text() for page in doc])
            doc.close()
            os.unlink(tmp_path)
            return text

        # ---------- DOCX ----------
        elif file_extension == ".docx":
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])

        # ---------- TXT ----------
        elif file_extension == ".txt":
            return str(uploaded_file.read(), "utf-8")

        # ---------- MARKDOWN ----------
        elif file_extension == ".md":
            return str(uploaded_file.read(), "utf-8")

        # ---------- EXCEL ----------
        elif file_extension in [".xlsx", ".xls"]:
            excel_data = pd.read_excel(uploaded_file, sheet_name=None)

            extracted_text = []
            for sheet_name, df in excel_data.items():
                extracted_text.append(f"\n--- Sheet: {sheet_name} ---\n")
                extracted_text.append(
                    df.fillna("").astype(str).to_string(index=False)
                )

            return "\n".join(extracted_text)

        return ""

    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {e}")
        return ""


# def extract_text_from_file(uploaded_file) -> str:
#     """Extracts text from PDF, TXT, DOCX, Excel, or Markdown."""
#     file_extension = Path(uploaded_file.name).suffix.lower()

#     try:
#         # ---------- PDF ----------
#         if file_extension == ".pdf":
#             with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
#                 tmp_file.write(uploaded_file.getvalue())
#                 tmp_path = tmp_file.name

#             doc = fitz.open(tmp_path)
#             text = "".join([page.get_text() for page in doc])
#             doc.close()
#             os.unlink(tmp_path)
#             return text

#         # ---------- DOCX ----------
#         elif file_extension == ".docx":
#             doc = Document(uploaded_file)
#             return "\n".join([para.text for para in doc.paragraphs])

#         # ---------- TXT ----------
#         elif file_extension == ".txt":
#             return str(uploaded_file.read(), "utf-8")

#         # ---------- MARKDOWN (.md) ----------
#         elif file_extension == ".md":
#             return str(uploaded_file.read(), "utf-8")

#         # ---------- EXCEL (.xlsx / .xls) ----------
#         elif file_extension in [".xlsx", ".xls"]:
#             excel_data = pd.read_excel(uploaded_file, sheet_name=None)

#             extracted_text = []
#             for sheet_name, df in excel_data.items():
#                 extracted_text.append(f"\n--- Sheet: {sheet_name} ---\n")

#                 # Convert sheet rows into text
#                 extracted_text.append(
#                     df.fillna("").astype(str).to_string(index=False)
#                 )

#             return "\n".join(extracted_text)

#         return ""

#     except Exception as e:
#         st.error(f"Error processing {uploaded_file.name}: {e}")
#         return ""



def generate_descriptions(document_text: str, feature_list: list, api_key: str, model_name: str, mode: str) -> dict:
    """Uses Groq to extract requirements based on mode."""
    
    if mode == "Structured":
        prompt = f"""
           You are a professional QA analyst. Your task is to create detailed descriptions of software features based on a Functional Specification Document (FSD). These descriptions will be used to write comprehensive test cases.

    **Instructions:**
    1. Read the entire `DOCUMENT_TEXT` provided below.
    2. For each feature name in the `FEATURE_LIST`, locate the corresponding section in the `DOCUMENT_TEXT`.
    3. Synthesize a detailed, paragraph-style description for each feature. This description must include:
        - The primary goal or purpose of the feature.
        - The event/trigger for each feature. 
        - The main user/actor who interacts with it.
        - The basic workflow or sequence of actions the user performs.
        - The alternate flow or sequence of actions the user performs.
        - The pre-condition 
        - The post condition
        - Validations/rules mentioned.
        - list every single UI field, status, and data point mentioned in the document. Mention all the descriptions and points covered under the feature in the document.
        - Do not summarize or condense any part of the document. Capture 100% of the details without omission.
    4. IMPORTANT: Use ONLY the information provided in the `DOCUMENT_TEXT`. Do not add any information or make assumptions.
    5. Format your final output as a single JSON object, where the keys are the exact feature names from the `FEATURE_LIST` and the values are the detailed string descriptions.
    6. Before generating the feature description, please add the following paragraph at the start of each description:
        Generate the most comprehensive and detailed set of test cases for the following feature description. Each line in each workflow must be tested through multiple functional, alternate, negative, boundary, UI, API, and access-control scenarios to ensure complete coverage.
        Below is the feature description:
        **FEATURE_LIST:** {json.dumps(feature_list)}
        **DOCUMENT_TEXT:** {document_text}
        """
    else:
        # Improved Unstructured Mode Prompt
        prompt = f"""
        You are a Master QA Requirements Engineer. Your goal is to perform a "Deep Extraction" of every single requirement, rule, and data point across multiple documents.

        **CRITICAL INSTRUCTIONS:**
        1. **Zero Omission:** Do not summarize. If the document mentions a specific validation, a character limit, a button color, or a background process, it MUST be included.
        2. **Logical Grouping:** Organize the extracted information into high-level Section Headings.
        3. **Comprehensive Detail:** For every requirement found, include:
            - The specific rule or logic.
            - Any data fields or formats mentioned.
            - Dependencies (what must happen before/after).
            - Error handling or "negative" paths described in the text.
        4. **Multiple Files:** The text provided contains content from several files. Synthesize them so that related requirements from different files are grouped under the same heading.
        5. **Formatting:** Return a valid JSON object. Each key is a Section Heading, and each value is a comprehensive, multi-paragraph string containing the detailed requirements.

        **DOCUMENT_TEXT (Combined from all files):**
        ---
        {document_text}
        ---

        **OUTPUT:**
        Respond with ONLY a valid JSON object.
        """

    try:
        client = Groq(api_key=api_key)

        with st.spinner("🤖 AI is analyzing your document and generating descriptions..."):
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "You are an expert QA analyst. Return ONLY valid JSON. No markdown, no extra text."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=4096
            )

        content = response.choices[0].message.content.strip()

        # Method 1: Try to extract JSON block if wrapped in ```
        import re
        # Extract JSON block first
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
        else:
            json_str = content

        # Safe parsing
        feature_descriptions = safe_json_loads(json_str)

        if not feature_descriptions:
            return {}

        return feature_descriptions
    except Exception as e: 
        st.error(f"Error communicating with Groq model: {e}") 
        return {}

def safe_json_loads(raw_text: str):
    """
    Attempts to repair common LLM JSON issues:
    - unescaped newlines
    - control characters
    - stray markdown
    """
    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        pass

    # -------- FIX 1: remove control chars --------
    cleaned = re.sub(r'[\x00-\x1F\x7F]', '', raw_text)

    # -------- FIX 2: remove markdown wrappers --------
    cleaned = cleaned.replace("```json", "").replace("```", "")

    # -------- FIX 3: normalize newlines --------
    cleaned = cleaned.replace('\r', '\\r').replace('\n', '\\n')

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        st.error(f"JSON parsing failed even after repair: {e}")
        st.code(cleaned[:1500] + "..." if len(cleaned) > 1500 else cleaned)
        return {}


# --- UI STYLING ---
# --- UI STYLING ---
# --- UI STYLING ---
st.markdown("""
<style>

/* ================== BUTTONS ================== */
div.stButton > button,
div.stDownloadButton > button,
button[kind="secondary"] {
    background-color: #a3d8ff !important;
    color: #000000 !important;
    border: none;
    border-radius: 8px;
    font-weight: 600;
}

/* Hover */
div.stButton > button:hover,
div.stDownloadButton > button:hover,
button[kind="secondary"]:hover {
    background-color: #8ccfff !important;
}

/* Disabled */
div.stButton > button:disabled {
    background-color: #d9eefc !important;
    color: #7a7a7a !important;
}

/* ================== FILE UPLOADER (Browse files) ================== */
button[data-testid="stFileUploaderBrowseButton"] {
    background-color: #a3d8ff !important;
    color: #000000 !important;
    border-radius: 8px;
    font-weight: 600;
}

button[data-testid="stFileUploaderBrowseButton"]:hover {
    background-color: #8ccfff !important;
}

/* ================== RADIO BUTTONS ================== */

/* Remove any background highlight on label */
.stRadio label {
    background-color: transparent !important;
}

/* Radio outer circle */
.stRadio input[type="radio"] {
    accent-color: #a3d8ff;
}

/* Prevent row highlight on hover */
.stRadio div[role="radiogroup"] label:hover {
    background-color: transparent !important;
}

/* ================== HEADERS ================== */
.main-header {
    font-size: 2.5rem;
    font-weight: bold;
    text-align: center;
}

.feature-box {
    background-color: #f0f8ff;
    padding: 1.5rem;
    border-radius: 10px;
    margin-bottom: 1rem;
    border-left: 4px solid #a3d8ff;
}

</style>
""", unsafe_allow_html=True)



# --- SESSION STATE & CONFIG ---

api_key = st.secrets.get("groq_api_key", None)
model_name = st.secrets.get("groq_default_model", "llama-3.3-70b-versatile")

if 'descriptions' not in st.session_state: st.session_state.descriptions = None
if 'combined_text' not in st.session_state: st.session_state.combined_text = None

# --- SIDEBAR ---
with st.sidebar:
    # st.header("⚙️ Configuration")
    
    # NEW: Mode Selection
    analysis_mode = st.radio(
        "Select Analysis Mode",
        ["Structured", "Unstructured"],
        help="Structured: Map to specific features. Unstructured: Extract everything discovered."
    )
    
    st.divider()
    
    # API Key check
    # if api_key: st.success("✅ API Key Loaded")
    # else: st.error("❌ API Key Missing")

    # Feature List (Only if Structured)
    features_list = []
    if analysis_mode == "Structured":
        st.subheader("📋 Features to Analyze")
        feature_text = st.text_area("Enter features (one per line)", placeholder="Login\nRegistration")
        if feature_text:
            features_list = [f.strip() for f in feature_text.split('\n') if f.strip()]
    else:
        st.info("ℹ️ In Unstructured mode, the AI will automatically identify all headings and requirements.")

# --- MAIN UI ---
st.markdown('<p class="main-header">📄 Requirement Analysis Tool</p>', unsafe_allow_html=True)

col1, col2 = st.columns([1, 1])

with col1:
    st.header("📤 Upload Documents")
    uploaded_files = st.file_uploader(
    "Upload FSD Files", 
    type=['pdf', 'txt', 'docx', 'xlsx', 'xls', 'md', 'png', 'jpg', 'jpeg'],
    accept_multiple_files=True
)
    
    if uploaded_files:
        if st.button("🔍 Process and Extract Text"):
            all_text = ""
            for uploaded_file in uploaded_files:
                with st.spinner(f"Reading {uploaded_file.name}..."):
                    all_text += f"\n--- Start of {uploaded_file.name} ---\n"
                    all_text += extract_text_from_file(uploaded_file)
            
            st.session_state.combined_text = all_text
            st.success(f"✅ Extracted text from {len(uploaded_files)} files.")

with col2:
    st.header("🚀 Run Analysis")
    
    # Logic for enabling button
    ready = api_key and st.session_state.combined_text
    if analysis_mode == "Structured" and not features_list:
        ready = False

    if st.button("✨ Generate Requirements", disabled=not ready):
        results = generate_descriptions(
            st.session_state.combined_text,
            features_list,
            api_key,
            model_name,
            analysis_mode
        )
        if results:
            st.session_state.descriptions = results

# --- RESULTS ---
if st.session_state.descriptions:
    st.header(f"📊 Results: {analysis_mode} Mode")
    
    st.download_button(
        "💾 Download Full Results (JSON)",
        data=json.dumps(st.session_state.descriptions, indent=2),
        file_name="requirements_output.json"
    )

    for idx, (title, content) in enumerate(st.session_state.descriptions.items(), 1):
        with st.container():
            st.markdown(f'<div class="feature-box"><h3>{idx}. {title}</h3></div>', unsafe_allow_html=True)
            st.markdown(content)
            st.divider()

# st.markdown("<p style='text-align: center; color: #666;'>Built with Streamlit & Groq AI</p>", unsafe_allow_html=True)







